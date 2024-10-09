from streamlit_chat import message as msg
import streamlit as st
import docx #pip install python-docx
import openai
import os
import io


openai.api_key = os.getenv("SENHA_OPEN_AI")

st.title("ChatGPT")
st.write("***")

if "hst_conversa" not in st.session_state:
    st.session_state.hst_conversa = []

pergunta = st.text_area("Digite sua pergunta: ")
btn = st.button("Enviar pergunta: ")

if btn:
    st.session_state.hst_conversa.append({"role": "user", "content": pergunta})
    retorno_openai = openai.ChatCompletion.create(
            model = "gpt-3.5-turbo",
            messages = st.session_state.hst_conversa,
            max_tokens = 500,
            n=1
    )
    st.write(retorno_openai["choices"][0]["message"]["content"])
    st.session_state.hst_conversa.append({"role": "assistant", "content": retorno_openai["choices"][0]["message"]["content"]})

if len(st.session_state.hst_conversa)>0:
    btn_salvar = st.button("salvar conteudo")

    if btn_salvar:
        trabalho = io.BytesIO()
        documento = docx.Document()
        documento.add_heading("Conteudo gerado ",level=1)
        for i in range(len(st.session_state.hst_conversa)):
                if i % 2 == 0:
                        documento.add_heading("Pergunta",level=2)
                        documento.add_paragraph(st.session_state.hst_conversa[i]["content"])
                else:
                        documento.add_heading("Resposta",level=2)
                        documento.add_paragraph(st.session_state.hst_conversa[i]["content"])
        documento.save(trabalho)
        st.download_button(label="Clique aqui para baixar o conteudo",data=trabalho,file_name="",mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    for i in range(len(st.session_state.hst_conversa)):
        if i % 2 == 0:
                msg("Voce: ",st.session_state.hst_conversa[i]["content"],is_user=True)
        else:
                msg("Resposta IA: ",st.session_state.hst_conversa[i]["content"])